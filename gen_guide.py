# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()
s = doc.sections[0]
s.page_width = Cm(21); s.page_height = Cm(29.7)
s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)
st = doc.styles[chr(78)+chr(111)+chr(114)+chr(109)+chr(97)+chr(108)]
st.font.name = chr(83)+chr(105)+chr(109)+chr(83)+chr(117)+chr(110)
st.font.size = Pt(12); st.paragraph_format.line_spacing = 1.5
for i in range(1,4):
    hs = doc.styles[fchr(39)Heading {i}chr(39)]
    hs.font.name = chr(83)+chr(105)+chr(109)+chr(72)+chr(101)+chr(105)
    hs.font.color.rgb = RGBColor(0,0,0); hs.font.bold = True

def P(text,b=False,c=False,ind=True,sz=12):
    p=doc.add_paragraph(); r=p.add_run(text)
    r.font.name=chr(83)+chr(105)+chr(109)+chr(83)+chr(117)+chr(110); r.font.size=Pt(sz); r.bold=b
    if c: p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not ind: p.paragraph_format.first_line_indent = Cm(0)

for _ in range(6): doc.add_paragraph()
tp = doc.add_paragraph(); tp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run(chr(22522)+chr(20110)+chr(32)+chr(68)+chr(101)+chr(101)+chr(112)+chr(67)+chr(104)+chr(97)+chr(116)+chr(32)+chr(65)+chr(73)+chr(32)+chr(43)+chr(32)+chr(68)+chr(101)+chr(101)+chr(112)+chr(83)+chr(101)+chr(101)+chr(107)+chr(32)+chr(22823)+chr(27169)+chr(22411)+chr(10)+chr(35770)+chr(25991)+chr(26234)+chr(33021)+chr(25490)+chr(29256)+chr(31995)+chr(32479)+chr(32)+chr(45)+chr(32)+chr(23436)+chr(25972)+chr(24320)+chr(21457)+chr(25805)+chr(20316)+chr(27493)+chr(39588))
r.font.name=chr(83)+chr(105)+chr(109)+chr(72)+chr(101)+chr(105); r.font.size=Pt(22); r.bold=True

OUT = rchr(39)C:\\Users\\xyh\\Desktop\\Test_Guide.docxchr(39)
doc.save(OUT)
print(fchr(39)OK: {os.path.getsize(OUT)} byteschr(39))
