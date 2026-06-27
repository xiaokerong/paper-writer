# -*- coding: utf-8 -*-
"""Complete guide generation - all content inline"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()
s = doc.sections[0]
s.page_width = Cm(21); s.page_height = Cm(29.7)
s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

st = doc.styles['Normal']
st.font.name = 'SimSun'; st.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
st.font.size = Pt(12); st.paragraph_format.line_spacing = 1.5; st.paragraph_format.first_line_indent = Cm(0.74)

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'SimHei'; hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    hs.font.color.rgb = RGBColor(0, 0, 0); hs.font.bold = True
    if i == 1: hs.font.size = Pt(18); hs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif i == 2: hs.font.size = Pt(15)
    else: hs.font.size = Pt(13)

def P(text, b=False, c=False, indent=True):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.name = 'SimSun'; r.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    r.font.size = Pt(12); r.bold = b
    if c: p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not indent: p.paragraph_format.first_line_indent = Cm(0)

def Code(text):
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0); p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text); r.font.name = 'Courier New'; r.font.size = Pt(9)

def B(text, lev=0):
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.74 + lev * 0.74)
    r = p.add_run(text); r.font.name = 'SimSun'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun'); r.font.size = Pt(12)

def N():
    doc.add_page_break()

def H1(t): doc.add_heading(t, level=1)
def H2(t): doc.add_heading(t, level=2)
def H3(t): doc.add_heading(t, level=3)

# ==========================================
# COVER PAGE
# ==========================================
for _ in range(5): doc.add_paragraph()
P('基于 DeepChat AI + DeepSeek 大模型\n论文智能排版系统 - 完整开发操作步骤', b=True, c=True, indent=False)
P('零基础 AI 辅助编程 全流程文档', c=True, indent=False)
N()

# ==========================================
# PART 1: AI 环境搭建
# ==========================================
H1('第一部分: AI 环境搭建')
H2('1.1 下载并安装 DeepChat AI 客户端')
P('打开浏览器, 访问 https://deepchatai.cn/ 进入 DeepChat 官网首页.')
P('在首页找到下载入口, 根据操作系统选择对应版本 (本文以 Windows 版本为例).')
P('点击下载按钮, 等待安装包下载完成.')
P('双击安装程序 (DeepChat-Setup.exe), 按安装向导提示完成安装, 保持默认选项即可.')
P('安装完成后桌面会出现 DeepChat 快捷方式图标.')
P('双击 DeepChat 图标启动客户端, 左侧为对话列表, 右侧为对话区域.')

H2('1.2 配置 DeepSeek 大模型服务')
P('步骤一: 进入 DeepChat 设置界面. 点击左下角齿轮图标进入设置面板.')
P('步骤二: 在设置面板中找到服务商设置选项卡, 找到 DeepSeek 选项并切换为启用.')
P('步骤三: 注册并获取 DeepSeek API Key.')
P('在浏览器打开 https://platform.deepseek.com/ 注册开发者账号.')
P('登录后在控制台左侧菜单找到 API Keys, 点击创建 API Key.')
P('复制生成的 sk- 开头的密钥 (注意: API Key 仅显示一次, 务必妥善保存).')
P('步骤四: 回到 DeepChat 设置面板, 将 API Key 粘贴到 DeepSeek 的输入框中.')
P('点击测试连接按钮确认连接成功.')
P('步骤五: 在 DeepChat 主界面顶部的模型选择下拉框中, 选择 deepseek-v4-pro 模型.')
P('该模型具备 128K 上下文窗口, 能够处理长文本, 非常适合论文相关任务.')

H2('1.3 DeepSeek API 计费说明')
P('DeepSeek API 采用按量付费模式, 注册时赠送一定额度免费试用金.')
P('开发过程中每次调用 API 都会消耗 tokens, 建议关注用量避免超出预算.')
P('在 DeepSeek 控制台的使用统计页面可实时查看 token 消耗情况.')

H2('1.4 我使用的提示词示例')
P('在 DeepChat 中输入以下提示来验证配置是否正常:')
Code('你好, 请简单介绍一下你自己和你的能力范围.')
P('如果 AI 正常回复, 说明 DeepSeek API 配置成功.')
N()

# ==========================================
# PART 2: 项目创建与需求分析
# ==========================================
H1('第二部分: 项目创建与需求分析')
H2('2.1 用 AI 规划项目')
P('在 DeepChat 中输入以下提示:')
Code('请帮我规划一个基于大模型的论文自动排版系统项目. 技术要求:')
Code('前端: React + TypeScript + Vite + Tailwind CSS')
Code('后端: Python FastAPI')
Code('AI: DeepSeek API')
Code('核心功能: Word上传 -> AI识别结构 -> 模板排版 -> 导出规范Word')
Code('请给出完整项目目录结构和各模块职责说明.')
P('AI 返回详细架构方案后, 确认方案合理, 开始搭建项目骨架.')

H2('2.2 创建项目文件夹')
P('在本地选择一个合适目录, 创建 paper-writer 文件夹作为项目根目录.')
P('在 DeepChat 中创建多个对话工作区, 按模块拆分:')
B('对话1: 项目架构 + 后端 main.py 开发')
B('对话2: 前端 App.tsx + 组件开发')
B('对话3: 智能排版引擎开发 (format_engine.py)')
B('对话4: 辅助功能模块开发 (选题/文献/查重等)')
B('对话5: 测试调试与问题修复')
P('每个对话有独立上下文, 避免内容过多导致模型遗忘.')

H2('2.3 项目初始化提示词')
P('我在 DeepChat 中使用的提示词:')
Code('请帮我初始化 paper-writer 项目. 前端使用 Vite+React+TS+Tailwind,')
Code('后端使用 FastAPI+python-docx+httpx. 请给出完整的初始化命令.')
N()

# ==========================================
# PART 3: 前端开发
# ==========================================
H1('第三部分: 前端开发')
H2('3.1 初始化前端项目')
P('在终端中执行:')
Code('cd paper-writer')
Code('npm create vite@latest frontend/paper-writer-app -- --template react-ts')
Code('cd frontend/paper-writer-app && npm install')
Code('npm install tailwindcss @tailwindcss/vite')

H2('3.2 配置 Tailwind CSS')
P('在 DeepChat 中:')
Code('请为我的 Vite React 项目配置 Tailwind CSS. 需要暗色科技风主题,')
Code('背景深蓝色 (#0a1628), 强调色青绿色 (#00d4aa).')
P('AI 返回配置后, 修改 vite.config.ts 和 index.css 文件.')

H2('3.3 创建核心文件 (通过 DeepChat 逐个生成)')
P('提示词示例:')
Code('请帮我定义论文排版系统的 TypeScript 类型接口:')
Code('PaperData (论文), GenerateRequest (生成请求), StyleOverrides (样式覆盖),')
Code('TemplateInfo (模板信息). 输出完整的 types.ts 文件.')

P('生成 api.ts:')
Code('请封装所有后端 API 调用函数, 使用 fetch, 统一错误处理.')
Code('包括: generatePaper, polishText, exportDocx, detectPaperStructure,')
Code('exportFormattedDocx, formatPipeline, recommendTopics, searchLiterature,')
Code('generateReview, generateProposal, analyzeOriginality 等.')

P('生成 App.tsx (分段):')
Code('请构建 App.tsx 的智能排版模块界面.')
Code('要求: 文件拖拽上传区 (仅.docx), 模板选择器, 当前模板参数展示,')
Code('导出Word按钮. 使用暗色科技风主题 (bg #0a1628, accent #00d4aa).')

H2('3.4 App.tsx 主界面架构')
P('App.tsx 是最大文件 (约1200+行), 包含:')
B('modules 数组: 7个模块定义 (智能排版/论文生成/润色/选题/文献/开题/查重)')
B('activeModule 状态: 控制当前显示模块')
B('侧边栏: 280px 宽, 模块导航按钮列表')
B('右侧内容区: 根据 activeModule 条件渲染对应模块UI')

H2('3.4.1 智能排版模块 UI (formatModule)')
B('文件拖拽区: HTML5 Drag & Drop API, accept=".docx"')
B('文件点击上传: 隐藏 input[type=file], onChange触发API调用')
B('模板选择器: 调用 GET /api/format/templates 获取模板列表')
B('样式覆盖面板: 自定义字体/字号/行距/页边距')
B('导出按钮: 调用 POST /api/format/export 下载 .docx')
B('状态变量: formatFile, detectedPaper, isDetecting, isExporting, selectedTemplate, formatError')

H2('3.5 遇到的问题与解决')
P('问题1: 文本输入和文件上传状态冲突.')
P('解决: textarea onChange 时清空 formatFile, 文件上传时清空 formatRawText, 两者互斥.')
P('问题2: 模板列表加载时机.')
P('解决: 在切换到 format 模块的 useEffect 中异步加载模板列表, 添加 templatesLoaded 标志防止重复请求.')
N()

# ==========================================
# PART 4: 后端开发
# ==========================================
H1('第四部分: 后端开发')
H2('4.1 初始化后端')
Code('mkdir backend && cd backend')
Code('python -m venv venv')
Code('venv\\Scripts\\activate')
Code('pip install fastapi uvicorn python-docx httpx python-dotenv pydantic')

H2('4.2 环境变量配置')
P('创建 .env 文件:')
Code('DEEPSEEK_API_KEY=sk-你的API密钥')
Code('DEEPSEEK_BASE_URL=https://api.deepseek.com/v1')

H2('4.3 main.py - FastAPI 入口')
P('提示词:')
Code('请用 FastAPI 创建后端应用. 配置 CORS 允许所有源. 应用名为论文智能写作与润色系统.')
Code('导入: FastAPI, HTTPException, CORS, StreamingResponse, UploadFile, BaseModel.')
Code('全局字典 papers_store 和 chat_history 用于内存存储.')

P('路由规划 - 31+ 个端点, 9大功能组:')
B('论文: POST /api/generate, /api/polish, /api/revise')
B('存取: GET /api/paper/{id}, /api/export/{id}')
B('排版: POST /api/format/detect, /export, /pipeline; GET /templates')
B('DOCX: POST /api/docx/parse, /to-text')
B('选题: POST /api/topic/recommend, /trend, /compare')
B('文献: POST /api/literature/search, /review')
B('开题: POST /api/proposal/generate')
B('查重: POST /api/plagiarism/analyze, /reduce')
B('可视化: POST /api/visual/wordcloud, /citations, /keyword-network, /hotspots')

H2('4.4 models.py - 数据模型')
P('提示词:')
Code('请定义所有API请求和响应的 Pydantic 模型类.')
Code('包括: PaperGenerateRequest, PolishRequest, MultiRoundRequest,')
Code('FormatDetectRequest, FormatExportRequest, TopicRequest,')
Code('LiteratureSearchRequest, ProposalRequest, PlagiarismAnalyzeRequest 等.')
P('每个模型定义字段类型、默认值和可选性, FastAPI 自动进行参数校验.')

H2('4.5 引擎模块 - paper_engine.py (论文生成)')
P('提示词:')
Code('请实现 paper_engine.py: 异步调用 DeepSeek API 生成论文.')
Code('函数: _call_deepseek(messages, temperature, max_tokens) - httpx 异步 POST')
Code('函数: _parse_json(content) - 解析AI返回的JSON, 包含4层容错')
Code('函数: generate_paper(topic, keywords, outline, ...) - 组装Prompt并调用API')
Code('函数: continue_writing(history, instruction, context) - 多轮修改')

P('_parse_json 容错机制:')
B('第1层: 直接 json.loads() 解析')
B('第2层: 修复非法转义序列 (LaTeX \\section 等命令)')
B('第3层: 正则提取最外层 JSON 对象 {.*}')
B('第4层: 自动闭合截断的 JSON (统计括号数量, 补全缺失的 } 和 ])')
B('额外: 移除 Markdown 代码块包裹 (```json...```)')

H2('4.6 引擎模块 - format_engine.py (智能排版核心)')
P('提示词:')
Code('请实现 format_engine.py: 论文结构自动识别与格式排版.')
Code('DETECT_SYSTEM_PROMPT: 结构识别专用System Prompt, 定义JSON输出格式和识别规则')
Code('detect_structure(raw_text): 调用DeepSeek API识别论文结构')
Code('load_template(template_name): 加载JSON模板 (gbt7713/general/course_paper)')
Code('_map_overrides(overrides): 扁平参数映射为模板嵌套结构')
Code('apply_template(paper_data, template, overrides): python-docx生成Word文档')

P('DETECT_SYSTEM_PROMPT 设计:')
B('定义完整JSON输出格式 (title/author/abstract/keywords/sections/conclusion/references)')
B('列举每种结构元素的识别规则 (标题=首行, 章节=编号开头段落, 参考文献=末段列表等)')
B('兜底规则: 无法识别时使用空值填充')
B('temperature=0.3 确保一致性, max_tokens=8192 防止截断')

H2('4.7 其他引擎模块')
P('提示词示例 (literature_engine.py):')
Code('请实现文献综述引擎. 集成 Semantic Scholar API 搜索论文.')
Code('函数: search_literature(topic, keywords, yearFrom, limit)')
Code('函数: generate_review(topic, papers) - 调用DeepSeek生成结构化综述')
B('polish_engine.py: 文本润色 (学术/简洁模式)')
B('topic_engine.py: 选题推荐 + 趋势/对比分析')
B('literature_engine.py: Semantic Scholar 搜索 + 综述生成')
B('proposal_engine.py: 开题报告自动生成')
B('plagiarism_engine.py: 原创性检测 + 智能降重')
B('chart_engine.py: 图表数据生成')
B('visual_engine.py: 词云/引文网络/热点图谱')
B('docx_parser.py: python-docx 提取 DOCX 纯文本')
B('export_engine.py: 生成规范 Word 文档')

H2('4.8 遇到的问题与解决')
P('问题1: AI 返回的 JSON 经常被 Markdown 代码块包裹.')
P('解决: 在 _parse_json 开头增加正则清理: re.sub(r\'^```(?:json)?\\s*\\n?\', \'\', content).')
P('问题2: 论文内容中的 LaTeX 命令 (\\section 等) 破坏 JSON 转义.')
P('解决: 用正则修复所有非法转义序列: re.sub(r\'\\\\(?![\"\\\\/bfnrtu])\', r\'\\\\\\\\\', content).')
P('问题3: 长论文导致 JSON 被截断.')
P('解决: 统计括号数量, 自动补全缺失的 } 和 ].')
P('问题4: 导入纯文本无法识别结构 (后来改为仅支持DOCX).')
P('解决: 移除文本框输入, 仅保留 DOCX 文件上传, 简化用户操作流程.')
N()

# ==========================================
# PART 5: 智能排版核心流程
# ==========================================
H1('第五部分: 智能排版核心功能开发')
H2('5.1 完整数据流 (12步)')
B('步骤1: 用户拖拽/点击上传 .docx 文件到前端')
B('步骤2: 前端 POST /api/format/pipeline (FormData)')
B('步骤3: 后端 docx_parser.extract_text_from_docx() 提取纯文本')
B('步骤4: 文本送入 format_engine.detect_structure()')
B('步骤5: detect_structure() 发送 System Prompt + 用户文本到 DeepSeek')
B('步骤6: DeepSeek 返回结构化 JSON')
B('步骤7: _parse_json() 解析 JSON (4层容错)')
B('步骤8: 返回 paper 数据 + 模板列表到前端')
B('步骤9: 前端渲染预览, 用户选择模板')
B('步骤10: 用户点击导出, POST /api/format/export')
B('步骤11: 后端 apply_template() 生成 DOCX')
B('步骤12: DOCX 以 StreamingResponse 返回, 浏览器触发下载')

H2('5.2 System Prompt 设计心得')
P('在 DeepChat 中反复调整 Prompt:')
Code('你是一个学术论文结构识别专家. 给定任意格式的纯文本,')
Code('请识别并提取论文的完整结构, 严格按JSON格式输出.')
Code('标题: 文本首行或题目/标题标签后的内容')
Code('作者: 标题下方独立行或作者/姓名标签后的内容')
Code('摘要: 以摘要/Abstract标记开始至关键词前的段落')
Code('章节: 以一/二/1./第X章等编号开头的段落及其后续内容')
Code('参考文献: 以参考文献/References开始直到文末的所有条目')
P('经过多次迭代, temperature 从 0.7 降到 0.3, max_tokens 从 4096 提升到 8192.')

H2('5.3 模板 JSON 结构')
P('每个模板 (gbt7713.json 等) 定义:')
B('page: 纸张大小, 上下左右边距 (cm)')
B('title: 字体 SimHei, 字号 16pt, 加粗居中')
B('body: 字体 SimSun, 字号 12pt, 行距 1.5, 首行缩进 0.74cm')
B('section_heading: 一级16pt / 二级14pt / 三级13pt, 均加粗')
B('reference: 字号 10.5pt, 悬挂缩进 0.74cm')
B('extras: 封面页/声明页/目录/致谢页开关')
N()

# ==========================================
# PART 6: 其他功能开发
# ==========================================
H1('第六部分: 其他功能模块开发')
H2('6.1 论文生成 (paper_engine.py)')
P('提示词:')
Code('请为论文生成功能设计 System Prompt. 引导模型输出包含以下部分的完整论文:')
Code('title/abstract/keywords/sections(引言/相关工作/方法/实验/结论)/references')
Code('支持传入真实引用文献, 在论文中适当位置引用.')
P('前端: 用户输入主题+关键词+大纲 -> 生成 -> 预览 -> 后续修改/润色/导出')

H2('6.2 文本润色 (polish_engine.py)')
P('提示词:')
Code('请实现文本润色功能. 支持学术风格 (正式严谨) 和简洁风格 (通俗精炼).')
Code('前端用 onMouseUp 捕获选中文本, 在浮动工具栏中展示润色结果.')

H2('6.3 选题推荐 (topic_engine.py)')
P('提示词:')
Code('请实现选题推荐引擎. 输入学科方向和学位级别, 调用DeepSeek分析研究热点.')
Code('输出: 推荐选题列表 (名称/创新点/可行性), 趋势数据, 对比分析.')

H2('6.4 文献综述 (literature_engine.py)')
P('提示词:')
Code('请集成 Semantic Scholar API 进行文献搜索.')
Code('搜索结果传给 DeepSeek, 生成包含研究现状/主要方法/存在问题/未来方向的结构化综述.')

H2('6.5 查重分析 (plagiarism_engine.py)')
P('提示词:')
Code('请实现查重引擎: 分析文本原创性, 识别可能抄袭的段落.')
Code('对高相似段落进行智能改写, 保持原意的同时变化表达.')

H2('6.6 图表与可视化 (chart_engine.py + visual_engine.py)')
P('提示词:')
Code('请生成适合 Recharts 渲染的图表数据 (折线图/雷达图/词云).')
P('前端使用 Recharts 库渲染, 生成 TrendLineChart, TopicRadarChart 等组件.')
N()

# ==========================================
# PART 7: 调试与测试
# ==========================================
H1('第七部分: 调试与测试')
H2('7.1 启动项目')
P('后端:')
Code('cd backend && python main.py    # http://localhost:8000')
P('前端:')
Code('cd frontend/paper-writer-app && npm run dev    # http://localhost:5173')
P('或双击 start-all.bat 一键启动.')

H2('7.2 常见问题排查')
P('问题1: DeepSeek API 返回空内容')
B('检查 .env 中 DEEPSEEK_API_KEY 是否正确')
B('检查 DeepSeek 账户余额 (platform.deepseek.com 控制台)')
B('检查网络能否访问 api.deepseek.com')
P('问题2: JSON 解析失败')
B('_parse_json 已增加 4 层容错 (直接解析/转义修复/正则提取/截断修复)')
B('如仍失败, 打印 content 前500字符定位具体问题')
P('问题3: 端口 8000 被占用')
B('打开任务管理器结束占用端口的 Python 进程')
B('或 netstat -ano | findstr :8000 找到PID后终止')
P('问题4: DOCX 解析失败')
B('确认是 .docx 格式 (非旧版 .doc)')
B('python-docx Document() 对象仅支持 Office Open XML 格式')

H2('7.3 测试清单')
B('GET /api/health -> status: ok')
B('POST /api/generate -> 验证 JSON 结构完整')
B('POST /api/format/pipeline -> 上传 .docx, 验证结构识别准确率')
B('POST /api/format/export -> 各模板导出, 检查排版规范性')
B('用户满意度问卷 -> 识别准确性/排版效果/操作便捷性/模板多样性/整体体验 (5分制)')

H2('7.4 调试经历分享')
P('在开发过程中遇到的主要问题和解决过程:')
B('智能排版 JSON 解析失败: 最初 DeepSeek 将 JSON 包裹在 Markdown 代码块中,')
B('  _parse_json 增加正则清理后解决. 后来又遇到 LaTeX 转义和截断问题,')
B('  逐步增加了 4 层容错. 问题通过 DeepChat 对话逐步排查和修复.')
B('PDF 导出编译失败: 尝试集成 LaTeX (xelatex) 生成 PDF,')
B('  但 xelatex 未安装、包依赖缺失、中文字体问题导致无法正常编译.')
B('  最终决定移除 PDF 导出, 专注于 DOCX 格式.')
B('TXT 文件识别失败: 最初支持上传 .txt 文件,')
B('  但 python-docx 无法解析 .txt 格式 (仅支持 ZIP/DOCX).')
B('  最终改为仅支持 DOCX 上传, 简化用户操作流程.')
N()

# ==========================================
# PART 8: PPT 与报告书
# ==========================================
H1('第八部分: 答辩 PPT 与项目报告书制作')
H2('8.1 答辩 PPT')
P('PPT 同样用 DeepChat 辅助生成:')
Code('请帮我设计毕业设计答辩PPT. 项目: 基于大模型的论文自动排版系统.')
Code('共10页: 封面/目录/研究背景/核心技术/系统架构/排版流程/')
Code('实验结果/功能全景/人员分工/总结致谢. 深色科技风主题.')
P('DeepChat 为每页生成 HTML 代码 (720pt x 405pt, 16:9),')
P('然后通过 html2pptx.js + PptxGenJS 转换为 PowerPoint 文件.')
P('实验结果页的柱状图用 PptxGenJS addChart API 在占位符位置自动生成.')
P('配色: 深蓝背景 #0a1628 + 青绿强调 #00d4aa.')

H2('8.2 项目报告书')
P('学校模板通过 python-docx 脚本自动填充:')
P('在 DeepChat 中上传模板文件, 解析段落结构, 生成填充脚本.')
P('报告内容包括: 研究背景/数据集/算法详解/系统架构/')
P('实验分析 (准确率94.8%, 满意度4.4/5.0)/人员分工/项目总结.')
N()

# ==========================================
# PART 9: 运行与部署
# ==========================================
H1('第九部分: 项目运行与部署')
H2('9.1 本地运行完整步骤')
Code('git clone .../paper-writer (或直接复制项目文件夹)')
Code('cd paper-writer/backend')
Code('python -m venv venv && venv\\Scripts\\activate')
Code('pip install fastapi uvicorn python-docx httpx python-dotenv pydantic')
Code('echo DEEPSEEK_API_KEY=sk-你的key > .env')
Code('python main.py')
Code('# 新终端:')
Code('cd paper-writer/frontend/paper-writer-app')
Code('npm install && npm run dev')
Code('# 浏览器访问 http://localhost:5173')

H2('9.2 技术栈总览')
B('前端: React 19 + TypeScript + Vite 6 + Tailwind CSS 4 + Recharts')
B('后端: Python FastAPI + uvicorn')
B('文档: python-docx')
B('AI: DeepSeek API (v4-pro, 128K上下文)')
B('学术: Semantic Scholar API')
B('PPT: html2pptx.js + PptxGenJS + Playwright')
B('开发: VS Code / PyCharm + Git')

H2('9.3 后续改进方向')
B('PDF导出: 集成 WeasyPrint 或 ReportLab')
B('更多模板: IEEE/ACM/Elsevier 期刊格式')
B('数学公式: LaTeX 公式解析')
B('排版预览: 下载前对比预览')
B('批量处理: 多文件同时排版')
B('在线编辑: WYSIWYG 编辑器')
B('用户系统: 注册/登录/历史记录')

N()

# ==========================================
# APPENDIX: 文件清单
# ==========================================
H1('附录: 项目文件清单')
Code('paper-writer/')
Code('  start-all.bat')
Code('  backend/')
Code('    main.py          FastAPI主入口 (31+路由)')
Code('    models.py        数据模型')
Code('    paper_engine.py  论文生成引擎')
Code('    format_engine.py 智能排版引擎 (核心)')
Code('    polish_engine.py 文本润色引擎')
Code('    topic_engine.py  选题推荐引擎')
Code('    literature_engine.py  文献综述引擎')
Code('    proposal_engine.py    开题报告引擎')
Code('    plagiarism_engine.py  查重分析引擎')
Code('    chart_engine.py       图表生成引擎')
Code('    visual_engine.py      可视化引擎')
Code('    docx_parser.py        DOCX文本提取')
Code('    export_engine.py      DOCX文档导出')
Code('    semantic_scholar.py   Semantic Scholar API')
Code('    templates/')
Code('      gbt7713.json general.json course_paper.json')
Code('    .env')
Code('  frontend/')
Code('    paper-writer-app/src/')
Code('      main.tsx, App.tsx, api.ts, types.ts, ChartComponents.tsx')

OUTPUT = r'C:\Users\xyh\Desktop\论文自动排版系统_开发操作步骤.docx'
doc.save(OUTPUT)
print(f'Done: {OUTPUT}')
print(f'Size: {os.path.getsize(OUTPUT)} bytes')
