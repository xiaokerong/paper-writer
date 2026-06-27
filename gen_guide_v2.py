# -*- coding: utf-8 -*-
"""Generate full operations guide DOCX"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os, json

doc = Document()
s = doc.sections[0]
s.page_width = Cm(21); s.page_height = Cm(29.7)
s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

st = doc.styles['Normal']
st.font.name = 'SimSun'; st.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
st.font.size = Pt(12); st.paragraph_format.line_spacing = 1.5; st.paragraph_format.first_line_indent = Cm(0.74)

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'SimHei'; hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    hs.font.color.rgb = RGBColor(0, 0, 0); hs.font.bold = True
    if i == 1: hs.font.size = Pt(18); hs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER; hs.paragraph_format.space_before = Pt(24); hs.paragraph_format.space_after = Pt(12)
    elif i == 2: hs.font.size = Pt(15); hs.paragraph_format.space_before = Pt(18); hs.paragraph_format.space_after = Pt(8)
    else: hs.font.size = Pt(13); hs.paragraph_format.space_before = Pt(12); hs.paragraph_format.space_after = Pt(6)

def P(text, b=False, c=False, ind=True):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.name = 'SimSun'; r.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun'); r.font.size = Pt(12); r.bold = b
    if c: p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not ind: p.paragraph_format.first_line_indent = Cm(0)

def Code(text):
    for line in text.split('\n'):
        p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0); p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line); r.font.name = 'Courier New'; r.font.size = Pt(9)

def B(text, lev=0):
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0); p.paragraph_format.left_indent = Cm(0.74 + lev * 0.74)
    r = p.add_run(text); r.font.name = 'SimSun'; r.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun'); r.font.size = Pt(12)

with open('guide_content.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

def render_section(sec):
    if 't' in sec:
        if sec.get('l') == 2:
            doc.add_heading(sec['t'], level=2)
        elif sec.get('l') == 3:
            doc.add_heading(sec['t'], level=3)
        else:
            doc.add_heading(sec['t'], level=1)
    if 'c' in sec:
        P(sec['c'])
    if 'code' in sec:
        Code(sec['code'])
    if 'b' in sec:
        for bl in sec['b']:
            if isinstance(bl, str):
                B(bl)
            elif isinstance(bl, dict):
                B(bl['t'], bl.get('l', 0))
    if 'page_break' in sec and sec['page_break']:
        doc.add_page_break()

# ---- TITLE PAGE ----
for _ in range(4):
    doc.add_paragraph()
P(data['cover_title'], b=True, c=True, ind=False)
P(data['cover_subtitle'], c=True, ind=False)
doc.add_page_break()

# ---- PARTS ----
for part_key, part_data in data.items():
    if not part_key.startswith('part'):
        continue
    if 'title' in part_data:
        doc.add_heading(part_data['title'], level=1)
    if 'sections' in part_data:
        for sec in part_data['sections']:
            render_section(sec)

OUTPUT = r'C:\Users\xyh\Desktop\论文自动排版系统_开发操作步骤.docx'
doc.save(OUTPUT)
print(f'OK: {OUTPUT}')
print(f'Size: {os.path.getsize(OUTPUT)} bytes')
