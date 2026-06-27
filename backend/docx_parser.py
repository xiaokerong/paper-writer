"""DOCX 解析器 — 从 Word 文件中提取结构化论文内容"""
import io
from docx import Document
from docx.oxml.ns import qn


def extract_text_from_docx(file_bytes: bytes) -> str:
    """从 DOCX 文件字节流中提取纯文本，保留段落结构

    解析策略：
    - 粗体短文本 → 可能是章节标题
    - 连续段落 → 正文
    - 编号列表 → 参考文献
    
    Returns:
        结构化的纯文本字符串
    """
    buffer = io.BytesIO(file_bytes)
    doc = Document(buffer)

    lines = []
    prev_was_heading = False

    for para in doc.paragraphs:
        text = _clean_text(para.text)
        if not text:
            continue

        # 判断是否为标题
        is_heading = _is_heading(para, text)

        if is_heading:
            if prev_was_heading:
                lines.append("")  # 标题间加空行
            lines.append(f"\n## {text}\n")
            prev_was_heading = True
        else:
            lines.append(text)
            prev_was_heading = False

    # 处理表格
    for table in doc.tables:
        lines.append("\n[表格]")
        for row in table.rows:
            cells = [_clean_text(cell.text) for cell in row.cells]
            lines.append(" | ".join(cells))
        lines.append("")

    return "\n".join(lines)


def extract_structured_paper(file_bytes: bytes) -> dict:
    """从 DOCX 文件中提取结构化论文数据（标题/作者/章节等）

    启发式规则 + python-docx API：
    - 首段居中/大号字体 → 标题
    - 第二段/小号字体居中对齐 → 作者
    - 以"摘要"/"Abstract"开头 → 摘要
    - 以编号("一、"、 "1."等)开头 → 章节标题
    - 以"参考文献"/"References"开头 → 参考文献区

    Returns:
        {"title": str, "author": str, "abstract": str, "sections": [...], "references": [...]}
    """
    buffer = io.BytesIO(file_bytes)
    doc = Document(buffer)

    result = {
        "title": "",
        "author": "",
        "abstract": "",
        "keywords": [],
        "sections": [],
        "conclusion": "",
        "references": [],
    }

    paragraphs = [p for p in doc.paragraphs if _clean_text(p.text)]

    if not paragraphs:
        return result

    # 阶段追踪
    phase = "header"  # header → abstract → body → references
    current_section = None
    current_content = []

    for i, para in enumerate(paragraphs):
        text = _clean_text(para.text)

        # 标题：第一段居中或大字体
        if phase == "header" and i == 0:
            result["title"] = text
            # 如果第二段短且字号小，可能是作者
            if len(paragraphs) > 1:
                p2_text = _clean_text(paragraphs[1].text)
                if len(p2_text) < 30 and _font_size_smaller(paragraphs[1], para):
                    result["author"] = p2_text
                    phase = "abstract"
                    continue
            phase = "abstract"
            continue

        # 作者识别
        if phase == "abstract" and len(text) < 30 and i <= 2:
            result["author"] = text
            continue

        # 摘要区
        if text.startswith("摘要") or text.startswith("Abstract"):
            abs_text = text.replace("摘要", "").replace("Abstract", "").strip("：: ")
            result["abstract"] = abs_text
            phase = "body"
            continue

        # 关键词
        if "关键词" in text or "Keywords" in text:
            kw_part = text.split("：")[-1] if "：" in text else text.split(":")[-1]
            result["keywords"] = [k.strip() for k in kw_part.replace("；", ";").split(";") if k.strip()]
            continue

        # 参考文献区
        if "参考文献" in text or "References" in text.replace(" ", ""):
            phase = "references"
            continue

        # 结论
        if text.replace(" ", "").startswith("结论") or text.replace(" ", "").startswith("结语"):
            phase = "body"
            if current_section:
                result["sections"].append({"heading": current_section, "content": "\n".join(current_content)})
            current_section = "结论"
            current_content = [text]
            continue

        # 致谢区
        if text.startswith("致谢") or text.startswith("致  谢"):
            phase = "references"  # 跳过致谢
            if current_section:
                result["sections"].append({"heading": current_section, "content": "\n".join(current_content)})
            current_section = None
            continue

        # 章节标题检测
        is_section_heading = _detect_section_heading(text)

        if is_section_heading and phase == "body":
            # 保存上一节
            if current_section:
                result["sections"].append({"heading": current_section, "content": "\n".join(current_content)})
            current_section = text
            current_content = []
            continue

        # 参考文献条目
        if phase == "references":
            if len(text) > 5:
                result["references"].append(text)
            continue

        # 正文
        if phase == "body" and current_section:
            current_content.append(text)

    # 保存最后一节
    if current_section and current_content:
        result["sections"].append({"heading": current_section, "content": "\n".join(current_content)})

    return result


def _clean_text(text: str) -> str:
    """清理文本：去除多余空白和不可见字符"""
    if not text:
        return ""
    return " ".join(text.split()).strip()


def _is_heading(para, text: str) -> bool:
    """判断段落是否为标题"""
    # 粗体 + 短文本
    is_bold = False
    if para.runs:
        first_run = para.runs[0]
        is_bold = first_run.bold or (first_run.font.size and first_run.font.size >= 180000)  # >= 14pt
    short = len(text) < 60
    return is_bold and short


def _font_size_smaller(para1, para2) -> bool:
    """判断 para1 的字号是否小于 para2"""
    try:
        sz1 = para1.runs[0].font.size if para1.runs else None
        sz2 = para2.runs[0].font.size if para2.runs else None
        if sz1 and sz2:
            return sz1 < sz2
    except Exception:
        pass
    return False


def _detect_section_heading(text: str) -> bool:
    """检测章节标题（中文编号、数字编号、英文编号）"""
    import re
    # 中文编号：一、二、三...
    if re.match(r'^[一二三四五六七八九十]+[、．.]', text):
        return True
    # 数字编号：1. 1.1 1.1.1
    if re.match(r'^\d+(\.\d+)*[\s.．、]', text):
        return True
    # 第X章/第X节
    if re.match(r'^第[一二三四五六七八九十\d]+[章节]', text):
        return True
    # Introduction/Methods/Results/Conclusion
    if re.match(r'^(I[VX]+\.?\s+|[A-Z][a-z]+\s*$)', text) and len(text) < 40:
        return True
    return False
