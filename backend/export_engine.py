"""DOCX 导出引擎 - 将论文导出为规范Word文档"""
import io
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def export_to_docx(paper: dict) -> bytes:
    """将论文数据导出为规范格式的Word文档"""
    doc = Document()
    
    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ── 样式设置 ──
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.first_line_indent = Cm(0.74)
    
    # 标题样式
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'SimHei'
        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        heading_style.font.color.rgb = RGBColor(0, 0, 0)
        heading_style.font.bold = True
        if i == 1:
            heading_style.font.size = Pt(16)
            heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif i == 2:
            heading_style.font.size = Pt(14)
        else:
            heading_style.font.size = Pt(13)

    # ── 论文标题 ──
    title = doc.add_heading(paper.get('title', '论文标题'), level=1)
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(18)

    # ── 摘要 ──
    doc.add_heading('摘  要', level=2)
    abstract = paper.get('abstract', '')
    p = doc.add_paragraph(abstract)
    p.paragraph_format.first_line_indent = Cm(0.74)

    # ── 关键词 ──
    keywords = paper.get('keywords', [])
    if keywords:
        p = doc.add_paragraph()
        run = p.add_run('关键词：')
        run.bold = True
        run.font.name = 'SimSun'
        run.font.size = Pt(12)
        p.add_run('；'.join(keywords))

    # ── 正文各章节 ──
    sections = paper.get('sections', [])
    for sec in sections:
        doc.add_heading(sec.get('heading', ''), level=2)
        content = sec.get('content', '')
        # 按段落分割
        paragraphs = content.split('\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                p = doc.add_paragraph(para_text)
                p.paragraph_format.first_line_indent = Cm(0.74)

    # ── 结论 ──
    doc.add_heading('结  论', level=2)
    conclusion = paper.get('conclusion', '')
    for para_text in conclusion.split('\n'):
        para_text = para_text.strip()
        if para_text:
            p = doc.add_paragraph(para_text)
            p.paragraph_format.first_line_indent = Cm(0.74)

    # ── 参考文献 ──
    doc.add_heading('参考文献', level=2)
    references = paper.get('references', [])
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph(f'[{i}] {ref}')
        p.paragraph_format.first_line_indent = Cm(0)

    # ── 保存到内存 ──
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
