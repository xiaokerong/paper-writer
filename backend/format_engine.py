"""智能排版引擎 — AI结构识别 + 模板应用"""
import json
import io
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from paper_engine import _call_deepseek, _parse_json

DETECT_SYSTEM_PROMPT = """你是学术论文结构识别专家。给定任意格式的纯文本，你需要识别并提取论文的完整结构。

请严格按以下JSON格式输出：
{
  "title": "论文标题",
  "author": "作者姓名（如未找到填"未知"）",
  "abstract": "摘要内容",
  "keywords": ["关键词1", "关键词2"],
  "sections": [
    {"heading": "章节标题", "content": "章节内容"}
  ],
  "conclusion": "结论内容",
  "references": ["参考文献1", "参考文献2"]
}

识别规则：
- 标题：文本首行或"题目"/"标题"标签后的内容
- 作者：标题下方独立行或"作者"/"姓名"标签后的内容
- 摘要：以"摘要"/"Abstract"标记开始至"关键词"前的段落
- 关键词：以"关键词"/"Keywords"标签后冒号/分号分隔的词
- 章节：以"一、"/"二、"/"1."/"第X章"等编号开头的段落及其后续内容
- 结论：以"结论"/"结语"/"总结"为标题的章节内容
- 参考文献：以"参考文献"/"References"开始直到文末的所有条目

如果某字段无法识别，用空值填充（字符串填""，数组填[]）。"""


TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates")

_SYSTEM_DEFAULTS = {
    "name": "系统默认",
    "description": "后备默认模板",
    "page": {"width_cm": 21, "height_cm": 29.7, "margin_top_cm": 2.54, "margin_bottom_cm": 2.54, "margin_left_cm": 3.17, "margin_right_cm": 3.17},
    "title": {"font": "SimHei", "font_eastasia": "SimHei", "size_pt": 16, "bold": True, "align": "center"},
    "author": {"font": "SimSun", "font_eastasia": "SimSun", "size_pt": 14, "align": "center", "show": True},
    "abstract": {"heading_font": "SimHei", "heading_size_pt": 14, "heading_bold": True, "body_font": "SimSun", "body_size_pt": 12},
    "keywords": {"font": "SimSun", "font_eastasia": "SimSun", "size_pt": 12},
    "body": {"font": "SimSun", "font_eastasia": "SimSun", "size_pt": 12, "line_spacing": 1.5, "first_line_indent_cm": 0.74},
    "section_heading": {"font": "SimHei", "font_eastasia": "SimHei", "level_1_size_pt": 16, "level_2_size_pt": 14, "level_3_size_pt": 13, "bold": True},
    "reference": {"font": "SimSun", "font_eastasia": "SimSun", "size_pt": 10.5, "format": "GB/T 7714", "hanging_indent_cm": 0.74},
    "extras": {},
}


async def detect_structure(raw_text: str) -> dict:
    """调用 DeepSeek 将纯文本解析为结构化论文数据"""
    if not raw_text.strip():
        raise ValueError("输入文本为空")

    content = await _call_deepseek(
        messages=[
            {"role": "system", "content": DETECT_SYSTEM_PROMPT},
            {"role": "user", "content": f"请识别以下论文的结构：\n\n{raw_text}"},
        ],
        temperature=0.3,
        max_tokens=8192,
    )
    return _parse_json(content)


def load_template(template_name: str) -> dict:
    """加载模板 JSON 文件，失败时回退系统默认"""
    file_path = os.path.join(TEMPLATES_DIR, f"{template_name}.json")
    if os.path.exists(file_path):
        with open(file_path, "r", encoding="utf-8") as f:
            return json.load(f)
    return dict(_SYSTEM_DEFAULTS)


def _deep_merge(base: dict, override: dict) -> dict:
    """深度合并，override 覆盖 base 中的对应键"""
    result = dict(base)
    for key, value in override.items():
        if key in result and isinstance(result[key], dict) and isinstance(value, dict):
            result[key] = _deep_merge(result[key], value)
        else:
            result[key] = value
    return result


def _map_overrides(overrides: dict) -> dict:
    """将前端扁平过载映射为模板嵌套结构"""
    mapping = {
        'bodyFont': ('body', 'font'),
        'bodySize': ('body', 'size_pt'),
        'lineSpacing': ('body', 'line_spacing'),
        'headingFont': ('section_heading', 'font'),
        'marginTop': ('page', 'margin_top_cm'),
        'marginBottom': ('page', 'margin_bottom_cm'),
        'marginLeft': ('page', 'margin_left_cm'),
        'marginRight': ('page', 'margin_right_cm'),
    }
    result = {}
    for flat_key, (section, key) in mapping.items():
        if flat_key in overrides and overrides[flat_key] is not None:
            if section not in result:
                result[section] = {}
            # Convert string number to float for numeric values
            val = overrides[flat_key]
            if isinstance(val, str) and key.endswith(('_cm', '_pt', 'line_spacing')):
                try:
                    val = float(val)
                except ValueError:
                    pass
            result[section][key] = val
    return result


def _save_extras(doc, paper, cfg):
    """Generate cover page, declaration, TOC, acknowledgment based on extras config"""
    extras = cfg.get("extras", {})
    
    # Cover page
    if extras.get("cover_page"):
        for _ in range(3):
            doc.add_paragraph("")
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(paper.get("title", ""))
        title_run.font.name = "SimHei"
        title_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        title_run.font.size = Pt(22)
        title_run.bold = True
        
        author = paper.get("author", "")
        if author:
            author_p = doc.add_paragraph()
            author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_p.add_run(author)
            author_run.font.name = "SimSun"
            author_run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            author_run.font.size = Pt(14)
        
        doc.add_page_break()
    
    # Declaration page
    if extras.get("declaration_page"):
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = title_p.add_run("原创性声明")
        tr.font.name = "SimHei"
        tr.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        tr.font.size = Pt(16)
        tr.bold = True
        
        decl = "本人郑重声明：所呈交的论文是本人在导师指导下独立完成的研究成果。除文中已注明引用的内容外，本论文不包含任何其他个人或集体已发表或撰写的研究成果。"
        p = doc.add_paragraph(decl)
        p.paragraph_format.first_line_indent = Cm(0.74)
        for run in p.runs:
            run.font.name = "SimSun"
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            run.font.size = Pt(12)
        
        doc.add_page_break()
    
    # TOC placeholder
    if extras.get("toc"):
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = title_p.add_run("目  录")
        tr.font.name = "SimHei"
        tr.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        tr.font.size = Pt(16)
        tr.bold = True
        
        doc.add_paragraph("（请在 Word 中右键更新域以自动生成完整目录）")
        doc.add_page_break()
    
    # Acknowledgment page
    if extras.get("acknowledgment_page"):
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = title_p.add_run("致  谢")
        tr.font.name = "SimHei"
        tr.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        tr.font.size = Pt(16)
        tr.bold = True
        
        ack = "本论文的完成离不开导师的悉心指导和各位同学的热心帮助，在此表示衷心的感谢。"
        p = doc.add_paragraph(ack)
        p.paragraph_format.first_line_indent = Cm(0.74)
        for run in p.runs:
            run.font.name = "SimSun"
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            run.font.size = Pt(12)


def apply_template(paper: dict, template: dict, overrides: dict | None = None) -> bytes:
    """按模板配置生成规范 Word 文档"""
    mapped_overrides = _map_overrides(overrides or {})
    effective = _deep_merge(template, mapped_overrides)

    doc = Document()

    # ── 页面设置 ──
    pg = effective["page"]
    section = doc.sections[0]
    section.page_width = Cm(pg["width_cm"])
    section.page_height = Cm(pg["height_cm"])
    section.top_margin = Cm(pg["margin_top_cm"])
    section.bottom_margin = Cm(pg["margin_bottom_cm"])
    section.left_margin = Cm(pg["margin_left_cm"])
    section.right_margin = Cm(pg["margin_right_cm"])

    # ── 正文默认样式 ──
    body_cfg = effective["body"]
    style = doc.styles['Normal']
    style.font.name = body_cfg["font"]
    style.element.rPr.rFonts.set(qn('w:eastAsia'), body_cfg.get("font_eastasia", body_cfg["font"]))
    style.font.size = Pt(body_cfg["size_pt"])
    style.paragraph_format.line_spacing = body_cfg["line_spacing"]
    style.paragraph_format.first_line_indent = Cm(body_cfg.get("first_line_indent_cm", 0.74))

    # ── 论文标题 ──
    t_cfg = effective["title"]
    title_p = doc.add_paragraph()
    title_p.alignment = _parse_align(t_cfg["align"])
    title_p.paragraph_format.space_after = Pt(12)
    run = title_p.add_run(paper.get("title", ""))
    _set_run_font(run, t_cfg)

    # ── 作者 ──
    a_cfg = effective["author"]
    if a_cfg.get("show", True) and paper.get("author"):
        author_p = doc.add_paragraph()
        author_p.alignment = _parse_align(a_cfg["align"])
        run = author_p.add_run(paper["author"])
        _set_run_font(run, a_cfg)

    # ── 摘要 ──
    ab_cfg = effective["abstract"]
    h_p = doc.add_paragraph()
    h_r = h_p.add_run("摘  要")
    h_r.bold = ab_cfg.get("heading_bold", True)
    h_r.font.name = ab_cfg["heading_font"]
    h_r.font.size = Pt(ab_cfg["heading_size_pt"])
    h_r.element.rPr.rFonts.set(qn('w:eastAsia'), ab_cfg["heading_font"])

    if paper.get("abstract"):
        ap = doc.add_paragraph(paper["abstract"])
        ap.paragraph_format.first_line_indent = Cm(0.74)
        for a_run in ap.runs:
            a_run.font.name = ab_cfg["body_font"]
            a_run.font.size = Pt(ab_cfg["body_size_pt"])

    # ── 关键词 ──
    keywords = paper.get("keywords", [])
    if keywords:
        kw_cfg = effective["keywords"]
        kp = doc.add_paragraph()
        b_r = kp.add_run("关键词：")
        b_r.bold = True
        _set_run_font(b_r, kw_cfg)
        k_r = kp.add_run("；".join(keywords))
        _set_run_font(k_r, kw_cfg)

    # ── 正文章节 ──
    sh = effective["section_heading"]
    sections = paper.get("sections", [])
    for sec in sections:
        heading = doc.add_paragraph()
        hd_r = heading.add_run(sec.get("heading", ""))
        hd_r.bold = sh.get("bold", True)
        hd_r.font.name = sh["font"]
        hd_r.font.size = Pt(sh["level_1_size_pt"])
        hd_r.element.rPr.rFonts.set(qn('w:eastAsia'), sh.get("font_eastasia", sh["font"]))

        content = sec.get("content", "")
        for para_text in content.split("\n"):
            para_text = para_text.strip()
            if para_text:
                p = doc.add_paragraph(para_text)
                p.paragraph_format.first_line_indent = Cm(body_cfg.get("first_line_indent_cm", 0.74))

    # ── 结论 ──
    if paper.get("conclusion"):
        ch = doc.add_paragraph()
        ch_r = ch.add_run("结  论")
        ch_r.bold = sh.get("bold", True)
        ch_r.font.name = sh["font"]
        ch_r.font.size = Pt(sh["level_1_size_pt"])
        ch_r.element.rPr.rFonts.set(qn('w:eastAsia'), sh.get("font_eastasia", sh["font"]))

        for para_text in paper["conclusion"].split("\n"):
            para_text = para_text.strip()
            if para_text:
                p = doc.add_paragraph(para_text)
                p.paragraph_format.first_line_indent = Cm(body_cfg.get("first_line_indent_cm", 0.74))

    # ── 参考文献 ──
    ref_cfg = effective["reference"]
    rh = doc.add_paragraph()
    rh_r = rh.add_run("参考文献")
    rh_r.bold = True
    rh_r.font.name = ref_cfg["font"]
    rh_r.font.size = Pt(ref_cfg.get("size_pt", 10.5))

    references = paper.get("references", [])
    for i, ref in enumerate(references, 1):
        rp = doc.add_paragraph(f"[{i}] {ref}")
        rp.paragraph_format.first_line_indent = Cm(0)
        for rp_run in rp.runs:
            rp_run.font.name = ref_cfg["font"]
            rp_run.font.size = Pt(ref_cfg.get("size_pt", 10.5))

    # ── 附加页面（封面、声明、目录、致谢）──
    _save_extras(doc, paper, effective)

    # ── 保存到内存 ──
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _set_run_font(run, cfg: dict):
    """设置 run 的字体属性"""
    run.font.name = cfg["font"]
    if cfg.get("font_eastasia"):
        run.element.rPr.rFonts.set(qn('w:eastAsia'), cfg["font_eastasia"])
    run.font.size = Pt(cfg["size_pt"])
    if cfg.get("bold"):
        run.bold = True


def _parse_align(align: str):
    """解析对齐方式字符串"""
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    return mapping.get(align, WD_ALIGN_PARAGRAPH.LEFT)
